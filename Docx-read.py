#!pip install python-docx
#!c:\python39\python.exe -m pip install --upgrade pip

from docx import Document
import os

directory=input('enter the file directory')
doc = Document(directory)
print(' header paragraps text ',doc.sections[0].header.paragraphs[0].text)

for x in range(0,5):
    print('text',doc.paragraphs[x].text)
